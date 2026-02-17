"""
Script para extrair conteúdo de arquivo DOCX
"""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx não está instalado")
    print("Execute: pip install python-docx")
    sys.exit(1)

def extract_docx_content(file_path):
    """Extrai todo o texto de um arquivo DOCX"""
    try:
        doc = Document(file_path)
        
        print("=" * 80)
        print(f"CONTEÚDO DO ARQUIVO: {Path(file_path).name}")
        print("=" * 80)
        print()
        
        for i, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                print(f"{para.text}")
        
        # Também extrair tabelas se houver
        if doc.tables:
            print("\n" + "=" * 80)
            print("TABELAS ENCONTRADAS:")
            print("=" * 80)
            
            for table_idx, table in enumerate(doc.tables, 1):
                print(f"\n--- Tabela {table_idx} ---")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    print(" | ".join(cells))
        
        print("\n" + "=" * 80)
        print("FIM DO ARQUIVO")
        print("=" * 80)
        
    except Exception as e:
        print(f"ERRO ao ler arquivo: {e}")
        sys.exit(1)

if __name__ == "__main__":
    file_path = r"C:\Users\SR APOIO\OneDrive\Documents\Projetos IA\agilepulse-dashboard\taredas squad.docx"
    extract_docx_content(file_path)
